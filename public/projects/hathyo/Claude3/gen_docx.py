#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo ClaudeSolution.docx cho Hathyo MLAM Claude3
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Trang & font mặc định ──────────────────────────────────────────────────
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ── Helpers ────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1e, 0x3a, 0x8f)
DKBLUE = RGBColor(0x0f, 0x17, 0x2a)
GREEN  = RGBColor(0x14, 0x6a, 0x30)
ORANGE = RGBColor(0xC0, 0x5A, 0x00)
RED    = RGBColor(0xA0, 0x00, 0x00)
TEAL   = RGBColor(0x00, 0x6A, 0x6A)
GRAY   = RGBColor(0x44, 0x44, 0x44)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if color:
            run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.bold = True
    return p

def add_para(text, bold_parts=None, indent=False, color=None, size=Pt(12)):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = size
        if color: run.font.color.rgb = color
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = size
        if color: run.font.color.rgb = color
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.8)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.8)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    return p

def add_note(text, note_type='info'):
    colors = {'info': ('1e3a8f', BLUE), 'warn': ('78350f', ORANGE), 'success': ('14532d', GREEN)}
    bg, tc = colors.get(note_type, colors['info'])
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = tc
    run.italic = True
    return p

def make_table(headers, rows, col_widths=None, header_bg='1e3a8f'):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if isinstance(cell_text, tuple):
                txt, bold, color = cell_text[0], cell_text[1] if len(cell_text)>1 else False, cell_text[2] if len(cell_text)>2 else None
            else:
                txt, bold, color = str(cell_text), False, None
            run = p.add_run(txt)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = bold
            if color: run.font.color.rgb = color
    # Col widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_space(n=1):
    for _ in range(n):
        doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HATHYO')
r.font.name = 'Times New Roman'; r.font.size = Pt(32); r.bold = True; r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HỆ THỐNG TIẾP THỊ LIÊN KẾT ĐA TẦNG (MLAM)')
r.font.name = 'Times New Roman'; r.font.size = Pt(18); r.bold = True; r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Multi-Level Affiliate Marketing Network')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.italic = True; r.font.color.rgb = GRAY

add_space(1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tài liệu Nghiệp vụ & Kỹ thuật – Phiên bản Claude3')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tháng 5/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(13)

add_space(1)
make_table(
    ['Thông tin dự án', 'Nội dung'],
    [
        ('Stack công nghệ', 'Java Spring Boot + PostgreSQL + ReactJS + React Native + AWS'),
        ('Thực hiện', 'Nguyễn Quốc Trọng + Nguyễn Minh Duy'),
        ('Pool hoa hồng', '10% (ac1=5%, ac2=2%, ac3=1.5%, ac4=1%, ac5=0.5%)'),
        ('Số cấp mạng lưới', 'Vô hạn (L0 → L1 → L2 → ... → Ln)'),
        ('Founders', '1.000 nút gốc (L0) – tạo bởi Admin'),
        ('Cơ chế đặc biệt', 'alfco – ví ảo hoàn cho người mua cấp <5'),
    ],
    col_widths=[5, 11]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# MỤC LỤC
# ══════════════════════════════════════════════════════════════════════════
add_heading('MỤC LỤC', 1, BLUE)
toc_items = [
    ('1.', 'Tổng Quan Hệ Thống Hathyo Hiện Hữu', '5'),
    ('1.1', 'Hạ tầng AWS (Bill tháng 3/2026)', '5'),
    ('1.2', 'Kiến trúc ứng dụng', '5'),
    ('1.3', 'Dịch vụ tích hợp', '6'),
    ('1.4', 'Cấu trúc phân quyền admin.hathyo.com', '6'),
    ('2.', 'Mô Hình MLAM – Vô Hạn Cấp, 5 Slot AC [MỚI]', '7'),
    ('2.1', 'Bối cảnh & So sánh thị trường', '7'),
    ('2.2', 'Nguyên tắc thiết kế (cập nhật Claude3)', '7'),
    ('2.3', 'Cấu trúc 5 Slot Hoa Hồng (ac1–ac5)', '8'),
    ('2.4', 'Phân tích tính hợp lý Pool 10%', '8'),
    ('3.', 'Hệ Thống 1.000 Founders (Nút Gốc)', '9'),
    ('3.1', 'Định nghĩa và Phân loại', '9'),
    ('3.2', 'Đặc quyền Founder', '9'),
    ('3.3', 'Mô phỏng tăng trưởng – Số Dunbar k=16', '10'),
    ('4.', 'Quy Tắc Nghiệp Vụ Chi Tiết [MỚI]', '11'),
    ('4.1', 'Tính toán commission_base', '11'),
    ('4.2', 'Quy tắc phân bổ AC – Công thức tổng quát', '11'),
    ('4.3', 'Các ví dụ phân bổ theo cấp người mua', '12'),
    ('4.4', 'Bảng alfco đầy đủ theo cấp người mua', '13'),
    ('4.5', 'Các trường hợp ngoại lệ (Edge Cases)', '13'),
    ('4.6', 'Timeline xử lý HH & alfco', '14'),
    ('5.', 'Thuế TNCN & Khung Tuân Thủ Pháp Lý', '15'),
    ('5.1', 'Cơ sở pháp lý', '15'),
    ('5.2', 'Nghĩa vụ khấu trừ thuế của sàn Hathyo', '15'),
    ('5.3', 'Ví dụ tính thuế thực tế', '16'),
    ('5.4', 'Quy trình hàng tháng', '16'),
    ('6.', 'Thiết Kế Cơ Sở Dữ Liệu PostgreSQL [MỚI]', '17'),
    ('6.1–6.8', 'Các bảng: affiliate_nodes, closure, commission_config, transactions, alfco_wallet, wallets, pit_declarations, payout_requests', '17'),
    ('7.', 'API Design – Java Spring Boot [MỚI]', '21'),
    ('7.1', 'Affiliate Member APIs', '21'),
    ('7.2', 'Registration with Referral Code', '22'),
    ('7.3', 'Admin APIs', '22'),
    ('8.', 'Commission Engine – Thuật Toán & Triển Khai [MỚI]', '23'),
    ('8.1', 'Tổng quan luồng xử lý', '23'),
    ('8.2', 'Pseudocode Java Spring Boot', '23'),
    ('8.3', 'Áp dụng alfco vào đơn hàng kế tiếp', '25'),
    ('8.4', 'Daily/Monthly Jobs', '25'),
    ('9.', 'Hướng Dẫn Triển Khai – Sprint Plan', '26'),
    ('9.1', 'Lộ trình tổng thể (8 tuần)', '26'),
    ('9.2', 'Chi tiết Sprint 1 – Database & Core Services', '26'),
    ('9.3', 'Checklist Frontend', '27'),
    ('9.4', 'Lưu ý bảo mật', '28'),
    ('9.5', 'Tích hợp với hệ thống hiện có', '28'),
    ('10.', 'Phụ Lục – Từ Điển Nghiệp Vụ', '29'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 if '.' in num and len(num) > 2 else 0)
    r1 = p.add_run(f'{num}  {title}')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    if '.' not in num or num.count('.') == 0:
        r1.bold = True; r1.font.color.rgb = BLUE
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 1 – TỔNG QUAN
# ══════════════════════════════════════════════════════════════════════════
add_heading('1. Tổng Quan Hệ Thống Hathyo Hiện Hữu', 1, BLUE)
add_para('Sàn thương mại điện tử Hathyo (hathyo.com) được vận hành bởi Công ty TNHH Cuộc Sống Vui Khỏe, địa chỉ 82 Phan Đăng Lưu, TP.HCM. Hệ thống hiện chạy trên AWS với chi phí ~$224/tháng (≈5,9 triệu VND).')

add_heading('1.1 Hạ tầng AWS (Bill tháng 3/2026)', 2, BLUE)
make_table(
    ['Dịch vụ', 'Chi phí (USD)', 'Mô tả', 'Ghi chú'],
    [
        ('Amazon EC2', '$93.34', 'Backend Java Spring Boot containers', 'Review instance khi scale'),
        ('Elastic Load Balancing', '$40.23', 'SSL termination, health check, load balancing', ''),
        ('ECR (Container Registry)', '$31.70', 'Docker images: backend + worker', ''),
        ('Amazon VPC', '$28.15', 'Network isolation, NAT Gateway', 'Private subnet RDS'),
        ('Amazon RDS (PostgreSQL)', '$23.19', 'Database chính toàn hệ thống', 'Multi-AZ khi scale'),
        ('AWS WAF + API GW + R53', '$7.42', 'WAF, API management, DNS', ''),
        ('S3 + CloudWatch + KMS', '$0.16', 'Storage, monitoring, encryption', ''),
        (('TỔNG', True, BLUE), ('$224.19', True, BLUE), '≈5,9 triệu VND/tháng', 'Account: 600627353258'),
    ],
    col_widths=[4.5, 2.5, 6, 3]
)
add_space()

add_heading('1.2 Kiến trúc ứng dụng', 2, BLUE)
make_table(
    ['Thành phần', 'Công nghệ', 'URL', 'Chức năng', 'Trạng thái'],
    [
        ('Backend API', 'Java Spring Boot + PostgreSQL', 'api.hathyo.com', 'REST API toàn bộ nghiệp vụ', ('✅ Production', False, GREEN)),
        ('Web Landing (Buyer)', 'ReactJS', 'hathyo.com', 'Visitor → User, mua hàng', ('✅ Production', False, GREEN)),
        ('Web Admin', 'ReactJS', 'admin.hathyo.com', 'Admin, nhân viên, Seller', ('✅ Production', False, GREEN)),
        ('App Android (Buyer)', 'React Native', 'Google Play', 'Tương đương web landing', ('✅ Released', False, GREEN)),
        ('App iOS (Buyer)', 'React Native', 'App Store', 'Đang hoàn thiện', ('🔄 In Progress', False, ORANGE)),
        ('App Seller', 'React Native', 'Internal', 'Quản lý gian hàng, đơn hàng', ('🔄 In Progress', False, ORANGE)),
    ],
    col_widths=[3.5, 3.5, 3, 4, 2]
)
add_space()

add_heading('1.3 Dịch vụ tích hợp', 2, BLUE)
make_table(
    ['Dịch vụ', 'Nhà cung cấp', 'Mục đích', 'Trạng thái'],
    [
        ('Xác minh địa chỉ', 'Vietbando API', 'Xác minh địa chỉ buyer/seller theo chuẩn ĐVHC VN', ('✅', False, GREEN)),
        ('Giao vận', 'GiaoHangTietKiem + Ahamove', 'Phí ship, vận đơn, tracking', ('✅', False, GREEN)),
        ('Tin nhắn thương hiệu', 'ZaloOA (ưu tiên) + SMS', 'OTP, thông báo đơn hàng, HH, thuế', ('✅', False, GREEN)),
        ('Hóa đơn điện tử', 'MatBao', 'Phát hành HĐDT, liên thông thuế', ('✅', False, GREEN)),
        ('Thanh toán', 'VietinBank (đang tích hợp)', 'Card, QR, CK; hiện chỉ CoD', ('🔄 CoD OK', False, ORANGE)),
    ],
    col_widths=[3.5, 4.5, 6, 2]
)
add_space()

add_heading('1.4 Cấu trúc phân quyền admin.hathyo.com', 2, BLUE)
make_table(
    ['Role', 'Quyền truy cập', 'Modules chính'],
    [
        (('Admin Hệ thống', True), 'Full access', 'Dashboard, Người dùng, Sản phẩm, Đơn hàng, MLAM Affiliate (MỚI), Cài đặt'),
        ('Nhân viên nội bộ', 'Partial – theo phân công', 'Tùy cấu hình từng nhân viên'),
        ('Người bán (Seller)', 'Chỉ gian hàng mình', 'Sản phẩm, Đơn hàng, Khuyến mãi, Doanh thu shop'),
    ],
    col_widths=[4, 4, 8]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 2 – MÔ HÌNH MLAM
# ══════════════════════════════════════════════════════════════════════════
add_heading('2. Mô Hình MLAM – Vô Hạn Cấp, 5 Slot AC [MỚI]', 1, BLUE)

add_heading('2.1 Bối cảnh & So sánh thị trường', 2, BLUE)
make_table(
    ['Chỉ số', 'Shopee', 'TikTok Shop', 'Hathyo MLAM (Claude3)'],
    [
        ('Tổng phí người bán chịu', '~40%', '~46%', ('≤ 30%', True, GREEN)),
        ('Pool affiliate', 'Unknown', 'Unknown', ('10% (5 AC slot)', True, GREEN)),
        ('Số cấp mạng lưới', 'N/A', 'N/A', ('Vô hạn (L0→Ln)', True, GREEN)),
        ('Ví ảo người mua cấp thấp', 'Không', 'Không', ('alfco (cấp <5)', True, GREEN)),
        ('CAC ước tính', '200K–500K/khách', '150K–400K', '10K–50K (referral)'),
    ],
    col_widths=[5, 3, 3, 5]
)
add_space()

add_heading('2.2 Nguyên tắc thiết kế (cập nhật Claude3)', 2, BLUE)
principles = [
    'HH chỉ phát sinh từ giao dịch thực (order DELIVERED/COMPLETED) – không phải từ tuyển dụng.',
    'Vô hạn cấp: Mạng lưới phát triển L0 → L1 → L2 → ... → Ln, không giới hạn depth.',
    '5 slot AC tính từ người mua ngược lên: L(n-1) nhận ac1=5%, L(n-2) nhận ac2=2%, L(n-3) nhận ac3=1.5%, L(n-4) nhận ac4=1%, L(n-5) nhận ac5=0.5%. Tổng = 10%.',
    'alfco: Khi người mua ở cấp n<5 (không đủ 5 upline), phần pool dư alfco được hoàn vào ví ảo người mua sau khi đơn giao thành công.',
    '1.000 Founders là nút gốc (L0), không có parent, được tạo bởi Admin.',
    'Không còn self-purchase bonus (khác Claude2): người tự mua không nhận thêm HH. Chỉ nhận alfco nếu n<5.',
    'Thuế TNCN 10% khấu trừ tại nguồn khi chi HH ≥ 2 triệu/lần. alfco miễn thuế (là khoản tặng người mua).',
]
for i, p in enumerate(principles):
    add_bullet(f'{i+1}. {p}')
add_space()

add_heading('2.3 Cấu trúc 5 Slot Hoa Hồng (ac1–ac5)', 2, BLUE)
add_note('Ký hiệu: n = cấp của người mua · alfc = min(n, 5) · alfco = 10% − Σ(aci, i=1..alfc) → ví ảo khi n<5', 'info')
make_table(
    ['Slot', 'Người nhận', 'Tỷ lệ', 'HH/1tr đơn', 'Lý do'],
    [
        (('ac1', True, RED), 'L(n-1) – upline trực tiếp', ('5%', True), '50.000đ', 'Công trực tiếp đưa người mua vào hệ thống'),
        (('ac2', True, ORANGE), 'L(n-2) – cách 2 bậc', ('2%', True), '20.000đ', 'Xây dựng được L(n-1), quan hệ sâu hơn'),
        (('ac3', True, RGBColor(0xB8,0x86,0x00)), 'L(n-3) – cách 3 bậc', ('1.5%', True), '15.000đ', 'Nền tảng trung gian trong chuỗi'),
        (('ac4', True, GREEN), 'L(n-4) – cách 4 bậc', ('1%', True), '10.000đ', 'Đóng góp vào cấu trúc sâu hơn'),
        (('ac5', True, BLUE), 'L(n-5) – cách 5 bậc', ('0.5%', True), '5.000đ', 'Giới hạn cuối chuỗi HH'),
        (('sum_ac', True, DKBLUE), ('Tổng pool', True), ('10%', True, DKBLUE), ('100.000đ', True, DKBLUE), 'Toàn bộ pool phân bổ cho tối đa 5 upline'),
        (('alfco', True, TEAL), 'Ví ảo người mua (n<5)', ('biến động', False, TEAL), '0–50.000đ', 'alfco = 10% − Σ(aci, i=1..alfc)'),
    ],
    col_widths=[2, 4, 2, 2.5, 5.5]
)
add_space()

add_heading('2.4 Phân tích tính hợp lý Pool 10%', 2, BLUE)
make_table(
    ['Hạng mục', '% tổng 30%', 'Giá trị/1tr đơn', 'Ghi chú'],
    [
        ('MLAM affiliate pool (ac1..ac5)', '10%', '100.000đ', 'Phân bổ tối đa 5 upline + alfco'),
        ('VAT (thuế GTGT trên phí DV)', '~5–6%', '50–60.000đ', 'Nghĩa vụ pháp lý'),
        ('Chi phí vận hành (logistics, tech, CS)', '~8–9%', '80–90.000đ', 'AWS + nhân sự + ops'),
        ('Lợi nhuận ròng sàn', '~5–7%', '50–70.000đ', 'Bền vững dài hạn'),
        (('Tổng phí người bán chịu', True), ('≤ 30%', True, GREEN), ('≤ 300.000đ', True, GREEN), 'vs Shopee 40%, TikTok 46% ✓'),
    ],
    col_widths=[5.5, 2.5, 3, 5]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3 – FOUNDERS
# ══════════════════════════════════════════════════════════════════════════
add_heading('3. Hệ Thống 1.000 Founders (Nút Gốc)', 1, BLUE)

add_heading('3.1 Định nghĩa và Phân loại', 2, BLUE)
add_para('Founders là 1.000 tài khoản đặc biệt được Admin tạo thủ công, là nút gốc L0 (depth_from_root = 0), không có parent. Họ có mã giới thiệu riêng (HAT-F0001 đến HAT-F1000) để phát triển cây mạng lưới bên dưới.')
make_table(
    ['Nhóm', 'Số lượng', 'Đối tượng', 'Quyền lợi đặc biệt'],
    [
        ('Nội bộ (Internal)', '~20', 'Người góp công phát triển từ đầu', 'Cổ phần hoặc quyền lợi tương đương + HH không giới hạn'),
        ('Đại diện tỉnh thành', '~50', '1–2 người/tỉnh theo 34 ĐVHC mới', 'Ưu tiên phát triển mạng địa phương'),
        ('KOL / Đối tác / Đầu tư', '~930', 'KOL, KOC, nhà đầu tư, đối tác quốc tế', 'Bán slot, quy đổi đầu tư, phân phối tương đương'),
    ],
    col_widths=[3.5, 2, 4, 6.5]
)
add_space()

add_heading('3.2 Đặc quyền Founder', 2, BLUE)
privileges = [
    'Mã định danh: HAT-F0001 → HAT-F1000, badge 🏆 FOUNDER trên mọi màn hình.',
    'Nhận HH từ ngày 1: Nhận ac5=0.5% khi cháu ở L5 mua hàng; ac4, ac3 từ những cháu gần hơn trong đúng vị trí chuỗi.',
    'Không giới hạn HH/tháng: Founder không bị cap tháng (khác với regular member).',
    'Không bị downgrade: Trừ khi Admin thu hồi thủ công (is_founder không tự động mất).',
    'Hỗ trợ ưu tiên: Dedicated support channel, onboarding riêng.',
]
for pr in privileges:
    add_bullet(pr)
add_space()

add_heading('3.3 Mô phỏng tăng trưởng – Số Dunbar k=16', 2, BLUE)
add_para('Theo lý thuyết Dunbar: mỗi người duy trì ~16 mối quan hệ có ảnh hưởng thực sự. Hệ thống thiết kế với k=16 (trung bình mỗi nút phát triển 16 nút trực tiếp).')
make_table(
    ['Cấp', 'Số lý thuyết (k=16)', 'Số thực (50% activation)', 'Ghi chú'],
    [
        ('L0 – Founder', '1.000', '1.000', 'Admin tạo thủ công'),
        ('L1', '16.000', '8.000', ''),
        ('L2', '256.000', '40.000', ''),
        ('L3', '4.096.000', '200.000', 'Bắt đầu chạm ngưỡng thị trường VN'),
        ('L4', '65.536.000', '500.000', 'Plateau, thị trường bão hòa'),
        ('L5+', '∞ (lý thuyết)', '~1.000.000', 'Vô hạn cấp – mạng vẫn hoạt động'),
    ],
    col_widths=[3, 4, 4, 5]
)
add_note('Tổng dân số VN ~100M, người dùng internet ~78M → các cấp 3–5 sẽ đến điểm bão hòa sau 12–24 tháng. Hệ thống vô hạn cấp vẫn hoạt động tốt vì HH phụ thuộc vào GMV, không chỉ số lượng node.', 'info')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4 – QUY TẮC NGHIỆP VỤ
# ══════════════════════════════════════════════════════════════════════════
add_heading('4. Quy Tắc Nghiệp Vụ Chi Tiết [MỚI]', 1, BLUE)

add_heading('4.1 Tính toán commission_base', 2, BLUE)
add_code(
    'commission_base = order.subtotal − order.discount_amount\n'
    '-- KHÔNG bao gồm: phí vận chuyển, VAT\n'
    '-- Ví dụ: Đơn 1.200.000đ − voucher 200.000đ − ship 30.000đ\n'
    '-- commission_base = 1.000.000đ\n'
    '-- affiliate pool   = 1.000.000 × 10% = 100.000đ'
)
add_space()

add_heading('4.2 Quy tắc phân bổ AC – Công thức tổng quát', 2, BLUE)
add_note('Người mua ở cấp n → alfc = min(n, 5) → Phân bổ ac1..ac(alfc) cho L(n-1)..L(n-alfc). Nếu alfc < 5: alfco = 10% − Σ(aci, i=1..alfc) → hoàn vào ví ảo người mua sau khi đơn COMPLETED.', 'info')
add_space()

add_heading('4.3 Các ví dụ phân bổ theo cấp người mua (đơn 1.000.000đ)', 2, BLUE)
make_table(
    ['Người mua', 'alfc', 'L(n-1) ac1=5%', 'L(n-2) ac2=2%', 'L(n-3) ac3=1.5%', 'L(n-4) ac4=1%', 'L(n-5) ac5=0.5%', 'alfco ví ảo'],
    [
        ('L1', '1', '50.000đ', '—', '—', '—', '—', ('50.000đ', True, TEAL)),
        ('L2', '2', '50.000đ', '20.000đ', '—', '—', '—', ('30.000đ', True, TEAL)),
        ('L3', '3', '50.000đ', '20.000đ', '15.000đ', '—', '—', ('15.000đ', True, TEAL)),
        ('L4', '4', '50.000đ', '20.000đ', '15.000đ', '10.000đ', '—', ('5.000đ', True, TEAL)),
        (('L5+ (vd L10)', False, GREEN), ('5', True, GREEN), '50.000đ', '20.000đ', '15.000đ', '10.000đ', '5.000đ', ('0đ', False, GRAY)),
    ],
    col_widths=[2.5, 1.2, 2.2, 2.2, 2.4, 2.2, 2.4, 2.9]
)
add_space()

add_heading('4.4 Bảng alfco đầy đủ theo cấp người mua', 2, BLUE)
make_table(
    ['Cấp người mua', 'alfc', 'Upline tổng nhận', 'alfco %', 'alfco/đơn 1tr', 'Ghi chú'],
    [
        ('L1', '1', '5%', ('5%', True, TEAL), ('50.000đ', True, TEAL), 'Founder nhận ac1; 5% còn lại → ví ảo'),
        ('L2', '2', '7%', ('3%', True, TEAL), ('30.000đ', True, TEAL), 'L1 ac1, L0 ac2'),
        ('L3', '3', '8.5%', ('1.5%', True, TEAL), ('15.000đ', True, TEAL), 'L2 ac1, L1 ac2, L0 ac3'),
        ('L4', '4', '9.5%', ('0.5%', True, TEAL), ('5.000đ', True, TEAL), 'L3 ac1, L2 ac2, L1 ac3, L0 ac4'),
        (('L5+', False, GREEN), ('5', False, GREEN), ('10%', False, GREEN), ('0', False, GRAY), ('—', False, GRAY), 'Đủ 5 upline, pool phân bổ hết – không alfco'),
    ],
    col_widths=[3, 1.5, 3, 2, 3, 4.5]
)
add_space()

add_heading('4.5 Các trường hợp ngoại lệ (Edge Cases)', 2, BLUE)
make_table(
    ['#', 'Kịch bản', 'Xử lý', 'Field DB'],
    [
        ('1', 'Người mua không có affiliate node', 'Toàn bộ 10% → platform pool', 'skip_reason = NO_NODE'),
        ('2', 'Chain ngắn (buyer ở L2, chỉ 2 upline)', 'Phân bổ đến hết chain; alfco → ví ảo', 'alfc_actual < 5'),
        ('3', 'Thành viên bị suspend trong chain', 'Skip level đó; cấp trên vẫn nhận', 'skip_reason = ACCOUNT_SUSPENDED'),
        ('4', 'Đơn hàng bị hoàn trả (REFUNDED)', 'Reverse toàn bộ commission + alfco records', 'status = REVERSED'),
        ('5', 'Đơn test / internal', 'Return sớm, không tạo bất kỳ record nào', 'is_test_order = true'),
        ('6', 'Đơn CoD chưa xác nhận giao', 'HH chỉ tạo khi order.status = DELIVERED', 'Async event after DELIVERED'),
        ('7', 'alfco chưa dùng khi đơn bị hoàn', 'alfco_wallet.status → REVERSED, trừ số dư', 'alfco_status = REVERSED'),
        ('8', 'HH < 2.000.000đ/lần chi', 'Không khấu trừ TNCN cho lần đó', 'pit_amount = 0'),
    ],
    col_widths=[0.8, 4.5, 5.5, 5.2]
)
add_space()

add_heading('4.6 Timeline xử lý HH & alfco', 2, BLUE)
make_table(
    ['Bước', 'Sự kiện', 'Order Status', 'HH Status', 'Hành động'],
    [
        ('1', 'Buyer đặt hàng', 'PENDING', '—', 'Tạo order; chưa làm gì với HH'),
        ('2', 'Seller xác nhận', 'CONFIRMED', '—', '—'),
        ('3', 'Đang giao hàng', 'SHIPPING', '—', '—'),
        ('4', 'Giao hàng thành công', 'DELIVERED', 'PENDING (tạo)', 'Async job: tạo commission records + alfco_pending; cộng pending_balance'),
        ('5', 'Buyer xác nhận (hoặc auto 3 ngày)', 'COMPLETED', 'APPROVED', 'HH → APPROVED; alfco → AVAILABLE; cộng available_balance'),
        ('6', 'User yêu cầu rút / cuối tháng', 'COMPLETED', 'PAID', 'Admin duyệt → chuyển tiền; khấu trừ thuế TNCN'),
        ('7', 'Đơn bị hoàn trả', 'REFUNDED', 'REVERSED', 'Reverse job; trừ lại balance + alfco'),
    ],
    col_widths=[1.2, 4, 2.5, 2.5, 5.8]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 5 – THUẾ TNCN
# ══════════════════════════════════════════════════════════════════════════
add_heading('5. Thuế TNCN & Khung Tuân Thủ Pháp Lý', 1, BLUE)

add_heading('5.1 Cơ sở pháp lý', 2, BLUE)
legal_items = [
    'Luật Thuế TNCN 2007 (sửa đổi 2012, 2014)',
    'Thông tư 111/2013/TT-BTC: Hướng dẫn thuế TNCN – khấu trừ 10% tại nguồn khi chi HH ≥ 2 triệu/lần',
    'Nghị định 126/2020/NĐ-CP: Quản lý thuế',
    'Luật Thương mại điện tử 2025 (122/2025/QH15): Hiệu lực 01/07/2026 – sàn khấu trừ và kê khai thuế thay người bán',
    'Thông tư 40/2021/TT-BTC: Hướng dẫn thuế đối với hộ kinh doanh, cá nhân kinh doanh',
]
for item in legal_items:
    add_bullet(item)
add_space()

add_heading('5.2 Nghĩa vụ khấu trừ thuế của sàn Hathyo', 2, BLUE)
make_table(
    ['Quy định', 'Chi tiết'],
    [
        (('Tỷ lệ khấu trừ', True), '10% TNCN trên mỗi khoản HH chi trả (KHÔNG áp dụng cho alfco)'),
        (('Ngưỡng áp dụng', True), 'Chi trả ≥ 2.000.000đ/lần → bắt buộc khấu trừ'),
        (('Kỳ kê khai', True), 'Hàng tháng – Tờ khai 05/KK-TNCN, nộp trước ngày 20 tháng tiếp theo'),
        (('Nộp tiền thuế', True), 'Nộp vào KBNN cùng kỳ với kê khai, trước ngày 20 tháng tiếp theo'),
        (('Chứng từ cho thành viên', True), 'Biên lai khấu trừ TNCN (Mẫu 03/BL), cấp cuối năm hoặc khi yêu cầu'),
        (('alfco (ví ảo)', True, TEAL), ('MIỄN thuế TNCN – là khoản tặng người mua (cashback), không phải hoa hồng', False, TEAL)),
        (('Khai quyết toán năm', True), 'Tờ khai 05/QTT-TNCN, nộp trước 31/3 năm sau'),
    ],
    col_widths=[5, 11]
)
add_space()

add_heading('5.3 Ví dụ tính thuế thực tế', 2, BLUE)
make_table(
    ['Khoản HH chi trả', 'Thuế TNCN 10%', 'Thực nhận (net)', 'Hành động hệ thống'],
    [
        ('1.500.000đ', 'Không KT (<2tr)', '1.500.000đ', 'Chuyển full vào ví'),
        ('2.000.000đ', ('200.000đ', False, ORANGE), '1.800.000đ', 'Tạo pit_transaction; cộng ví net'),
        ('5.000.000đ', ('500.000đ', False, ORANGE), '4.500.000đ', 'Tạo pit_transaction; báo cáo GDT'),
        ('50.000.000đ', ('5.000.000đ', False, RED), '45.000.000đ', 'Bắt buộc KYC MST cá nhân trước khi chi'),
    ],
    col_widths=[3.5, 3, 3.5, 6]
)
add_space()

add_heading('5.4 Quy trình hàng tháng', 2, BLUE)
monthly_steps = [
    'Ngày 1–5 tháng M+1: Tổng hợp toàn bộ HH đã chi trả trong tháng M theo từng thành viên. alfco không đưa vào tờ khai.',
    'Ngày 6–10: Tạo tờ khai XML theo chuẩn GDT, review nội bộ.',
    'Ngày 11–19: Submit lên etax.gdt.gov.vn và nộp tiền thuế vào KBNN.',
    'Ngày 20 (hạn cuối): Hoàn tất nghĩa vụ khai và nộp thuế.',
    'Cuối năm: Tổng hợp quyết toán; cấp biên lai 03/BL cho từng thành viên.',
]
for s in monthly_steps:
    add_bullet(s)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 6 – DB SCHEMA
# ══════════════════════════════════════════════════════════════════════════
add_heading('6. Thiết Kế Cơ Sở Dữ Liệu PostgreSQL [MỚI]', 1, BLUE)
add_para('Các bảng mới cần tạo thêm vào schema hiện tại. Bảng users và orders hiện có sẽ được mở rộng thêm cột. Tổng 8 bảng mới + 2 ALTER TABLE.')

add_heading('6.1 Mở rộng bảng users (ALTER TABLE)', 2, BLUE)
add_code(
    'ALTER TABLE users\n'
    '  ADD COLUMN IF NOT EXISTS affiliate_node_id UUID REFERENCES affiliate_nodes(node_id),\n'
    '  ADD COLUMN IF NOT EXISTS is_affiliate_member BOOLEAN NOT NULL DEFAULT false;'
)

add_heading('6.2 Bảng affiliate_nodes – Nút trong mạng lưới', 2, BLUE)
add_code(
    'CREATE TABLE affiliate_nodes (\n'
    '  node_id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
    '  user_id          UUID NOT NULL UNIQUE REFERENCES users(user_id),\n'
    '  parent_node_id   UUID REFERENCES affiliate_nodes(node_id), -- NULL = Founder L0\n'
    '  referral_code    VARCHAR(20) NOT NULL UNIQUE,  -- HAT-F#### hoặc HAT-#-#####\n'
    '  is_founder       BOOLEAN NOT NULL DEFAULT false,\n'
    '  founder_category VARCHAR(20),  -- INTERNAL, PROVINCE, PARTNER, INTERNATIONAL\n'
    '  province_note    TEXT,\n'
    '  depth_from_root  INTEGER NOT NULL DEFAULT 0, -- 0=Founder, 1,2...N (VÔ HẠN)\n'
    '  status           VARCHAR(20) NOT NULL DEFAULT \'ACTIVE\',\n'
    '  joined_at        TIMESTAMPTZ NOT NULL DEFAULT NOW(),\n'
    '  created_by_admin UUID,\n'
    '  CONSTRAINT chk_founder_no_parent CHECK (NOT is_founder OR parent_node_id IS NULL)\n'
    ');\n'
    '-- Lưu ý: Bỏ CONSTRAINT chk_depth – vô hạn cấp, không giới hạn depth\n'
    'CREATE INDEX idx_anode_user   ON affiliate_nodes(user_id);\n'
    'CREATE INDEX idx_anode_parent ON affiliate_nodes(parent_node_id);\n'
    'CREATE INDEX idx_anode_code   ON affiliate_nodes(referral_code);\n'
    'CREATE INDEX idx_anode_depth  ON affiliate_nodes(depth_from_root);'
)

add_heading('6.3 Bảng affiliate_tree_closure – Closure Table', 2, BLUE)
add_code(
    'CREATE TABLE affiliate_tree_closure (\n'
    '  ancestor_id   UUID NOT NULL REFERENCES affiliate_nodes(node_id),\n'
    '  descendant_id UUID NOT NULL REFERENCES affiliate_nodes(node_id),\n'
    '  depth         INTEGER NOT NULL, -- 0=self, 1=parent, ... (không giới hạn)\n'
    '  PRIMARY KEY (ancestor_id, descendant_id)\n'
    ');\n'
    'CREATE INDEX idx_closure_desc ON affiliate_tree_closure(descendant_id, depth);\n'
    'CREATE INDEX idx_closure_anc  ON affiliate_tree_closure(ancestor_id, depth);\n\n'
    '-- Khi thêm node B (child) với parent A:\n'
    'INSERT INTO affiliate_tree_closure(ancestor_id, descendant_id, depth)\n'
    '  SELECT :b, :b, 0\n'
    '  UNION ALL\n'
    '  SELECT ancestor_id, :b, depth + 1\n'
    '    FROM affiliate_tree_closure WHERE descendant_id = :a;\n\n'
    '-- Query 5 upline gần nhất của buyer X:\n'
    'SELECT n.*, c.depth FROM affiliate_tree_closure c\n'
    'JOIN affiliate_nodes n ON n.node_id = c.ancestor_id\n'
    'WHERE c.descendant_id = :buyerNodeId AND c.depth BETWEEN 1 AND 5\n'
    'ORDER BY c.depth ASC;'
)

add_heading('6.4 Bảng commission_config – Cấu hình tỷ lệ AC', 2, BLUE)
add_code(
    'CREATE TABLE commission_config (\n'
    '  config_id      UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
    '  ac_index       SMALLINT NOT NULL,      -- 1..5 (ac1..ac5)\n'
    '  ac_code        VARCHAR(10) NOT NULL,   -- \'ac1\'..\'ac5\'\n'
    '  rate_percent   DECIMAL(5,2) NOT NULL,  -- 5.00, 2.00, 1.50, 1.00, 0.50\n'
    '  effective_from TIMESTAMPTZ NOT NULL DEFAULT NOW(),\n'
    '  effective_to   TIMESTAMPTZ,            -- NULL = hiện hành\n'
    '  created_by     UUID,\n'
    '  CONSTRAINT chk_ac_index CHECK (ac_index BETWEEN 1 AND 5)\n'
    ');\n'
    '-- Seed data mặc định (Claude3, sum_ac = 10%):\n'
    'INSERT INTO commission_config(ac_index, ac_code, rate_percent) VALUES\n'
    '  (1,\'ac1\',5.00),(2,\'ac2\',2.00),(3,\'ac3\',1.50),(4,\'ac4\',1.00),(5,\'ac5\',0.50);'
)

add_heading('6.5 Bảng commission_transactions – Giao dịch HH', 2, BLUE)
add_code(
    'CREATE TABLE commission_transactions (\n'
    '  tx_id               UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
    '  order_id            UUID NOT NULL REFERENCES orders(order_id),\n'
    '  buyer_node_id       UUID NOT NULL REFERENCES affiliate_nodes(node_id),\n'
    '  beneficiary_node_id UUID REFERENCES affiliate_nodes(node_id), -- NULL=platform pool\n'
    '  buyer_level         INTEGER NOT NULL,   -- depth_from_root của người mua\n'
    '  ac_index            SMALLINT,           -- 1..5; NULL nếu platform pool\n'
    '  distance            SMALLINT,           -- khoảng cách upline-buyer (=ac_index)\n'
    '  commission_base_vnd BIGINT NOT NULL,\n'
    '  rate_snapshot       DECIMAL(5,2) NOT NULL,\n'
    '  gross_amount_vnd    BIGINT NOT NULL,\n'
    '  pit_rate            DECIMAL(5,2) DEFAULT 10.00,\n'
    '  pit_amount_vnd      BIGINT NOT NULL DEFAULT 0,\n'
    '  net_amount_vnd      BIGINT NOT NULL,\n'
    '  status              VARCHAR(20) NOT NULL DEFAULT \'PENDING\',\n'
    '  skip_reason         VARCHAR(50),\n'
    '  period_month        CHAR(7) NOT NULL,   -- \'YYYY-MM\'\n'
    '  created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),\n'
    '  approved_at TIMESTAMPTZ, paid_at TIMESTAMPTZ, reversed_at TIMESTAMPTZ,\n'
    '  CONSTRAINT chk_status CHECK (status IN\n'
    '    (\'PENDING\',\'APPROVED\',\'PAID\',\'SKIPPED\',\'REVERSED\',\'PLATFORM_POOL\'))\n'
    ');'
)

add_heading('6.6 Bảng alfco_wallet – Ví ảo người mua [MỚI]', 2, BLUE)
add_code(
    'CREATE TABLE alfco_wallet (\n'
    '  alfco_id         UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
    '  order_id         UUID NOT NULL REFERENCES orders(order_id),\n'
    '  buyer_node_id    UUID NOT NULL REFERENCES affiliate_nodes(node_id),\n'
    '  buyer_level      INTEGER NOT NULL,\n'
    '  alfc_actual      SMALLINT NOT NULL,     -- min(buyer_level, 5)\n'
    '  alfco_percent    DECIMAL(5,2) NOT NULL, -- % dư = 10 - sum(aci, i=1..alfc)\n'
    '  alfco_amount_vnd BIGINT NOT NULL,\n'
    '  status           VARCHAR(20) NOT NULL DEFAULT \'PENDING\',\n'
    '  -- PENDING → AVAILABLE (sau COMPLETED) → USED → REVERSED\n'
    '  available_at     TIMESTAMPTZ,\n'
    '  used_order_id    UUID,\n'
    '  used_at          TIMESTAMPTZ,\n'
    '  reversed_at      TIMESTAMPTZ,\n'
    '  created_at       TIMESTAMPTZ NOT NULL DEFAULT NOW()\n'
    ');'
)

add_heading('6.7 Bảng commission_wallets – Ví HH', 2, BLUE)
add_code(
    'CREATE TABLE commission_wallets (\n'
    '  node_id              UUID PRIMARY KEY REFERENCES affiliate_nodes(node_id),\n'
    '  available_balance    BIGINT NOT NULL DEFAULT 0,\n'
    '  pending_balance      BIGINT NOT NULL DEFAULT 0,\n'
    '  total_earned_gross   BIGINT NOT NULL DEFAULT 0,\n'
    '  total_pit_withheld   BIGINT NOT NULL DEFAULT 0,\n'
    '  total_paid_out       BIGINT NOT NULL DEFAULT 0,\n'
    '  current_month_gross  BIGINT NOT NULL DEFAULT 0,\n'
    '  alfco_available      BIGINT NOT NULL DEFAULT 0,  -- ví ảo sẵn dùng\n'
    '  alfco_pending        BIGINT NOT NULL DEFAULT 0,\n'
    '  alfco_total_received BIGINT NOT NULL DEFAULT 0,\n'
    '  updated_at           TIMESTAMPTZ NOT NULL DEFAULT NOW()\n'
    ');'
)

add_heading('6.8 Bảng pit_declarations & payout_requests', 2, BLUE)
add_code(
    'CREATE TABLE pit_declarations (\n'
    '  declaration_id     UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
    '  period_month       CHAR(7) NOT NULL,\n'
    '  node_id            UUID NOT NULL REFERENCES affiliate_nodes(node_id),\n'
    '  total_gross_vnd    BIGINT NOT NULL,\n'
    '  total_pit_vnd      BIGINT NOT NULL,\n'
    '  total_net_vnd      BIGINT NOT NULL,\n'
    '  tx_count           INTEGER NOT NULL DEFAULT 0,\n'
    '  declaration_status VARCHAR(20) DEFAULT \'DRAFT\',\n'
    '  submitted_at       TIMESTAMPTZ,\n'
    '  gdt_reference      VARCHAR(100),\n'
    '  UNIQUE (period_month, node_id)\n'
    ');\n\n'
    'CREATE TABLE payout_requests (\n'
    '  payout_id        UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
    '  node_id          UUID NOT NULL REFERENCES affiliate_nodes(node_id),\n'
    '  requested_amount BIGINT NOT NULL,\n'
    '  pit_amount       BIGINT NOT NULL DEFAULT 0,\n'
    '  net_amount       BIGINT NOT NULL,\n'
    '  payment_method   VARCHAR(20) NOT NULL,  -- MOMO/ZALOPAY/VIETINBANK\n'
    '  payment_account  VARCHAR(50) NOT NULL,\n'
    '  bank_name        VARCHAR(100),\n'
    '  account_holder   VARCHAR(200),\n'
    '  status           VARCHAR(20) DEFAULT \'PENDING\',\n'
    '  admin_note       TEXT,\n'
    '  requested_at     TIMESTAMPTZ NOT NULL DEFAULT NOW(),\n'
    '  processed_at     TIMESTAMPTZ,\n'
    '  processed_by     UUID\n'
    ');'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 7 – API DESIGN
# ══════════════════════════════════════════════════════════════════════════
add_heading('7. API Design – Java Spring Boot [MỚI]', 1, BLUE)
add_para('Base URL: /api/v1 | Auth: Bearer JWT | Content-Type: application/json')

add_heading('7.1 Affiliate Member APIs (User đã đăng nhập)', 2, BLUE)
make_table(
    ['Method', 'Endpoint', 'Mô tả', 'Response chính'],
    [
        ('POST', '/affiliate/join', 'User tham gia MLAM (tạo node)', 'node_id, referral_code, depth'),
        ('GET', '/affiliate/me', 'Thông tin MLAM của tôi', 'node, depth, wallet, upline_count'),
        ('GET', '/affiliate/my-code', 'Referral code + QR + link', 'code, link, qr_data'),
        ('GET', '/affiliate/my-network', 'Stats mạng lưới theo cấp', 'per_level: {count, orders, gmv, my_ac}'),
        ('GET', '/affiliate/my-upline', 'Chuỗi upline đến Founder', '[{level, name, code, ac_slot}]'),
        ('GET', '/affiliate/my-commissions', 'Lịch sử HH (filter ac_index/status/month)', 'paginated, incl pit'),
        ('GET', '/affiliate/my-wallet', 'Số dư ví HH', 'available, pending, total_earned, pit'),
        ('GET', '/affiliate/my-alfco', 'Ví ảo alfco', 'alfco_available, alfco_pending, history'),
        ('POST', '/affiliate/request-payout', 'Yêu cầu rút tiền HH', 'payout_id, estimated_at'),
        ('GET', '/affiliate/tax/my-certificate', 'Biên lai khấu trừ TNCN', 'PDF stream'),
        ('GET', '/affiliate/ac-rates', 'Xem cấu hình AC hiện tại', '[{ac_index, rate, effective_from}]'),
    ],
    col_widths=[1.5, 5, 5, 4.5]
)
add_space()

add_heading('7.2 Registration with Referral Code', 2, BLUE)
add_code(
    'POST /api/v1/auth/register\n'
    '{\n'
    '  "phone": "0901234567",\n'
    '  "password": "...",\n'
    '  "full_name": "Nguyen Van A",\n'
    '  "referral_code": "HAT-3-00042"  // optional\n'
    '}\n'
    '→ Response:\n'
    '{\n'
    '  "user_id": "uuid",\n'
    '  "access_token": "...",\n'
    '  "affiliate_node_created": true,\n'
    '  "my_referral_code": "HAT-4-01234",\n'
    '  "depth_in_tree": 4,\n'
    '  "alfc_if_buy_now": 4,\n'
    '  "alfco_pct_if_buy_now": 0.5\n'
    '}'
)
add_space()

add_heading('7.3 Admin APIs (PLATFORM_ADMIN)', 2, BLUE)
make_table(
    ['Method', 'Endpoint', 'Mô tả'],
    [
        ('POST', '/admin/affiliate/founders', 'Tạo Founder mới (1.000 slot max)'),
        ('GET', '/admin/affiliate/founders', 'Danh sách Founders với filter'),
        ('PUT', '/admin/affiliate/founders/{id}/status', 'Suspend/Activate Founder'),
        ('GET', '/admin/affiliate/network/{nodeId}', 'Xem cây mạng lưới từ 1 node (max 5 cấp)'),
        ('GET', '/admin/affiliate/transactions', 'Tất cả commission transactions với filter'),
        ('POST', '/admin/affiliate/transactions/approve-batch', 'Duyệt nhiều transactions'),
        ('GET', '/admin/affiliate/alfco/summary', 'Tổng hợp alfco theo cấp và kỳ'),
        ('GET', '/admin/affiliate/payouts/pending', 'Danh sách payout chờ duyệt'),
        ('POST', '/admin/affiliate/payouts/{id}/approve', 'Duyệt payout'),
        ('PUT', '/admin/affiliate/commission-config', 'Cập nhật AC rates (validate sum = 10%)'),
        ('GET', '/admin/affiliate/tax/monthly-summary', 'Tổng hợp thuế TNCN theo tháng'),
        ('POST', '/admin/affiliate/tax/submit-gdt', 'Nộp kê khai lên GDT API'),
    ],
    col_widths=[1.5, 6.5, 8]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 8 – COMMISSION ENGINE
# ══════════════════════════════════════════════════════════════════════════
add_heading('8. Commission Engine – Thuật Toán & Triển Khai [MỚI]', 1, BLUE)

add_heading('8.1 Tổng quan luồng xử lý', 2, BLUE)
add_para('Commission Engine chạy ASYNC sau khi OrderDeliveredEvent được publish. Không block checkout flow. Spring @Async + @Transactional đảm bảo atomicity.')

add_heading('8.2 Pseudocode Java Spring Boot', 2, BLUE)
add_code(
    '@Service\n'
    'public class CommissionEngineService {\n\n'
    '    // AC_RATES[0]=ac1=5%, AC_RATES[4]=ac5=0.5%\n'
    '    private static final double[] AC_RATES = {5.0, 2.0, 1.5, 1.0, 0.5};\n'
    '    private static final double SUM_AC = 10.0;\n\n'
    '    @EventListener @Async("commissionPool")\n'
    '    public void onOrderDelivered(OrderDeliveredEvent e) {\n'
    '        processCommission(e.getOrderId());\n'
    '    }\n\n'
    '    @Transactional\n'
    '    public void processCommission(UUID orderId) {\n'
    '        Order order = orderRepo.findById(orderId).orElseThrow();\n'
    '        if (order.isTestOrder()) return;\n\n'
    '        long base = order.getSubtotal() - order.getDiscountAmount();\n'
    '        if (base <= 0) return;\n\n'
    '        Optional<AffiliateNode> buyerOpt = nodeRepo.findByUserId(order.getBuyerUserId());\n'
    '        if (buyerOpt.isEmpty()) {\n'
    '            platformPool.add(base * SUM_AC / 100, orderId, "NO_NODE"); return;\n'
    '        }\n'
    '        AffiliateNode buyer = buyerOpt.get();\n'
    '        int n = buyer.getDepthFromRoot();       // cấp của người mua\n'
    '        int alfc = Math.min(n, 5);              // số cấp nhận HH\n\n'
    '        // Query tối đa 5 upline gần nhất qua closure table\n'
    '        List<UplineInfo> ancestors = closureRepo.findAncestors(\n'
    '            buyer.getNodeId(), 1, alfc);\n\n'
    '        List<CommissionConfig> acConfigs = configService.getActiveConfigs();\n'
    '        long totalPaid = 0;\n\n'
    '        for (int i = 0; i < alfc; i++) {\n'
    '            if (i >= ancestors.size()) break;\n'
    '            UplineInfo up = ancestors.get(i);\n'
    '            if (up.getStatus() != ACTIVE) {\n'
    '                logSkip(orderId, up, "ACCOUNT_SUSPENDED"); continue;\n'
    '            }\n'
    '            double rate  = acConfigs.get(i).getRatePercent();\n'
    '            long gross   = (long)(base * rate / 100);\n'
    '            long pit     = gross >= 2_000_000 ? (long)(gross * 0.10) : 0;\n'
    '            long net     = gross - pit;\n'
    '            saveCommissionTx(orderId, buyer, up.getNode(), i+1, base, gross, pit, net);\n'
    '            walletService.addPending(up.getNode(), net);\n'
    '            totalPaid += gross;\n'
    '        }\n\n'
    '        // Tính alfco – ví ảo người mua\n'
    '        long poolTotal = (long)(base * SUM_AC / 100);\n'
    '        long alfco     = poolTotal - totalPaid;\n'
    '        if (alfco > 0) {\n'
    '            double alfcoPct = SUM_AC - sumAcRates(alfc);\n'
    '            saveAlfcoPending(orderId, buyer, n, alfc, alfcoPct, alfco);\n'
    '            walletService.addAlfcoPending(buyer, alfco);\n'
    '        }\n'
    '    }\n\n'
    '    private double sumAcRates(int alfc) {\n'
    '        double sum = 0;\n'
    '        for (int i = 0; i < alfc; i++) sum += AC_RATES[i];\n'
    '        return sum;\n'
    '    }\n\n'
    '    // Khi order COMPLETED: approve HH + alfco\n'
    '    @EventListener @Async("commissionPool")\n'
    '    public void onOrderCompleted(OrderCompletedEvent e) {\n'
    '        commissionRepo.approveByOrder(e.getOrderId());\n'
    '        alfcoRepo.approveByOrder(e.getOrderId());\n'
    '        walletService.movePendingToAvailable(e.getOrderId());\n'
    '    }\n'
    '}'
)
add_space()

add_heading('8.3 Daily/Monthly Jobs', 2, BLUE)
make_table(
    ['Job', 'Schedule', 'Mô tả'],
    [
        ('Commission Auto-Approve', 'Daily 03:00', 'Order DELIVERED >3 ngày → COMPLETED → HH PENDING → APPROVED; alfco → AVAILABLE'),
        ('Monthly PIT Summary', '1st/month 01:00', 'Tổng hợp PIT theo từng node tháng trước → INSERT pit_declarations'),
        ('Wallet Monthly Reset', '1st/month 01:01', 'Reset current_month_gross = 0 cho tất cả wallets'),
        ('Payout Processing', 'Daily 09:00', 'Xử lý payout đã APPROVED → gọi MoMo/VietinBank API; retry on fail'),
        ('alfco Expiry Notify', 'Weekly', 'alfco AVAILABLE >90 ngày chưa dùng → thông báo qua ZaloOA'),
    ],
    col_widths=[4, 3, 9]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 9 – SPRINT PLAN
# ══════════════════════════════════════════════════════════════════════════
add_heading('9. Hướng Dẫn Triển Khai – Sprint Plan', 1, BLUE)

add_heading('9.1 Lộ trình tổng thể (8 tuần)', 2, BLUE)
make_table(
    ['Sprint', 'Tuần', 'Focus', 'Deliverable'],
    [
        ('Sprint 1', '1–2', 'Database + Core Services', '8 bảng mới; ReferralTreeService; CommissionEngine; AlfcoService'),
        ('Sprint 2', '3–4', 'APIs + Founder Management', '25+ REST APIs; Admin Founder; Order event hook; PIT; alfco API'),
        ('Sprint 3', '5–6', 'Frontend Web (ReactJS)', 'affiliate-landing + affiliate-admin (Claude3); tích hợp admin.hathyo.com'),
        ('Sprint 4', '7–8', 'Mobile App + Go-Live', 'React Native: tab Affiliate + alfco + payout + deep link; soft launch 10% user'),
    ],
    col_widths=[2, 1.5, 4, 8.5]
)
add_space()

add_heading('9.2 Chi tiết Sprint 1 – Database & Core Services', 2, BLUE)
make_table(
    ['Task', 'Mô tả', 'Người làm', 'Estimate', 'Priority'],
    [
        ('DB-01', 'Migration SQL tạo 8 bảng mới (incl. alfco_wallet); seed commission_config AC1-AC5', 'Backend #1', '3 ngày', ('P0', True, RED)),
        ('DB-02', 'Tạo indexes; ENUM types; constraints; triggers trên closure table', 'Backend #1', '1 ngày', ('P0', True, RED)),
        ('BE-01', 'AffiliateNodeService: create node, gen referral code, calc depth', 'Backend #1', '2 ngày', ('P0', True, RED)),
        ('BE-02', 'ReferralTreeService: insert closure table khi join; findAncestors(nodeId, 1, 5)', 'Backend #2', '2 ngày', ('P0', True, RED)),
        ('BE-03', 'CommissionEngineService: AC distribution + PIT + alfco (full implementation)', 'Backend #2', '3 ngày', ('P0', True, RED)),
        ('BE-04', 'AlfcoService: pending/available/used/reversed; apply-alfco vào đơn mới', 'Backend #1', '2 ngày', ('P0', True, RED)),
        ('BE-05', 'WalletService: pending/available balance; PIT tracking; monthly reset job', 'Backend #1', '1 ngày', ('P0', True, RED)),
        ('QA-01', 'Unit test: tree insert, AC calc (n<5 + n≥5), alfco, edge cases, reverse', 'QA', '2 ngày', ('P0', True, RED)),
    ],
    col_widths=[1.5, 7, 2.5, 2, 1.5]
)
add_space()

add_heading('9.3 Checklist Frontend – Màn hình cần xây', 2, BLUE)
make_table(
    ['Màn hình', 'Platform', 'Priority'],
    [
        ('Affiliate Landing – tổng quan, AC rates, mạng lưới, alfco', 'Web + App', ('P0', True, RED)),
        ('Affiliate Admin – toàn bộ quản trị MLAM', 'Web admin', ('P0', True, RED)),
        ('Trang share referral code + QR', 'Web + App', ('P0', True, RED)),
        ('Dashboard HH + mạng lưới + AC breakdown theo cấp', 'Web + App', ('P0', True, RED)),
        ('Ví HH + Rút tiền (min 200K)', 'Web + App', ('P0', True, RED)),
        ('Ví ảo alfco + lịch sử nhận + áp dụng vào đơn', 'Web + App', ('P0', True, RED)),
        ('Thuế TNCN + Chứng từ 03/BL', 'Web + App', ('P0', True, RED)),
        ('Admin: Founder management (tạo, khóa, phân loại)', 'Web admin', ('P0', True, RED)),
        ('Admin: AC Config + bảng alfco simulation', 'Web admin', ('P0', True, RED)),
        ('Admin: Payout approval + PIT', 'Web admin', ('P0', True, RED)),
        ('Admin: Tax declarations + GDT submit', 'Web admin', ('P1', False, ORANGE)),
        ('Admin: Audit Log đầy đủ (alfco, AC config, payout)', 'Web admin', ('P1', False, ORANGE)),
    ],
    col_widths=[9, 3, 2]
)
add_space()

add_heading('9.4 Lưu ý bảo mật', 2, BLUE)
security = [
    'Referral code: dùng base62 nanoid (20 ký tự) để tránh brute-force đoán mã.',
    'Rate limiting: API validate referral code giới hạn 10 request/phút/IP.',
    'AC calculation: server-side hoàn toàn, không tin vào client.',
    'Payout OTP: Xác nhận qua ZaloOA/SMS trước khi tạo payout request.',
    'Dual-control: Payout >10 triệu cần 2 admin duyệt.',
    'Audit log: Mọi thay đổi AC config, Founder create, payout approve đều có audit trail.',
    'Anti-circular: Closure table constraint tự động ngăn A→B→A (circular referral).',
    'alfco fraud: Phát hiện khi alfco_total_received >50% tổng doanh thu của node → alert admin.',
    'Depth bombing: cfgMaxDepth mặc định 50 trong commission engine để tránh vòng lặp vô tận.',
]
for s in security:
    add_bullet(s)
add_space()

add_heading('9.5 Tích hợp với hệ thống hiện có', 2, BLUE)
make_table(
    ['Tích hợp', 'Mô tả', 'Cách thực hiện'],
    [
        ('MatBao (HĐDT)', 'HH của sàn → doanh thu → phát hành HĐDT qua MatBao', 'CommissionEngineService → trigger MatBao API khi platform pool confirmed'),
        ('VietinBank', 'Payout tự động qua VietinBank API', 'PayoutService gọi disbursement API; fallback manual nếu fail'),
        ('ZaloOA + SMS', 'Thông báo HH nhận, alfco available, payout, OTP rút tiền', 'NotificationService.sendCommissionAlert() + sendAlfcoAlert()'),
        ('PostgreSQL hiện có', 'Thêm 8 bảng mới; FK đến users và orders', 'Migration script Flyway; backward compatible'),
        ('React Web', 'Thêm Affiliate section vào admin.hathyo.com; standalone landing', 'New React module; lazy loaded; tham khảo Claude3 HTML files'),
        ('React Native App', 'Thêm tab Affiliate vào app buyer (screen stack mới)', 'New screen stack; deep link referral code'),
    ],
    col_widths=[3, 5.5, 7.5]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 10 – PHỤ LỤC
# ══════════════════════════════════════════════════════════════════════════
add_heading('10. Phụ Lục – Từ Điển Nghiệp Vụ', 1, BLUE)
terms = [
    ('MLAM', 'Multi-Level Affiliate Marketing – Tiếp Thị Liên Kết Đa Tầng. Hệ thống phân phối hoa hồng qua mạng lưới giới thiệu nhiều cấp không giới hạn.'),
    ('Founder (L0)', '1.000 nút gốc được Admin tạo; depth_from_root = 0; không có parent; tham gia nhận HH khi đứng đúng vị trí trong chuỗi 5 cấp.'),
    ('Cấp n (Ln)', 'Vị trí trong cây: L0=Founder, L1 do Founder giới thiệu, L2 do L1 giới thiệu, ... Ln không giới hạn.'),
    ('ac1..ac5', '5 slot hoa hồng tính từ người mua ngược lên: ac1=5% (gần nhất), ac2=2%, ac3=1.5%, ac4=1%, ac5=0.5% (xa nhất). Tổng sum_ac = 10%.'),
    ('alfc', 'Affiliate Levels for Commission = min(n, 5). Số cấp upline thực sự nhận HH khi người mua ở cấp n.'),
    ('alfco', 'Affiliate Level Commission Offset = 10% − Σ(aci, i=1..alfc). Phần pool dư khi n<5; hoàn vào ví ảo người mua sau khi đơn COMPLETED.'),
    ('Ví ảo (alfco wallet)', 'Số dư alfco của người mua. Dùng để giảm giá đơn hàng kế tiếp. Không rút được thành tiền mặt. Miễn thuế TNCN.'),
    ('commission_base', 'subtotal − discount_amount (không tính ship, VAT). Dùng để tính affiliate pool = 10% × commission_base.'),
    ('affiliate pool', '10% của commission_base. Tổng phân bổ cho tối đa 5 upline; phần còn lại (nếu n<5) → alfco ví ảo.'),
    ('platform_pool', 'Phần HH bị skip (không có affiliate node, account suspended) quay về sàn làm doanh thu.'),
    ('PIT', 'Personal Income Tax – Thuế Thu Nhập Cá Nhân. 10% khấu trừ tại nguồn khi chi HH ≥ 2 triệu/lần.'),
    ('Closure Table', 'Bảng lưu tất cả quan hệ ancestor-descendant để query upline chain trong O(1), không cần CTE đệ quy mỗi đơn hàng.'),
    ('pending_balance', 'HH đang chờ order COMPLETED; chưa rút được.'),
    ('available_balance', 'HH đã APPROVED; sẵn sàng rút (tối thiểu 200.000đ).'),
    ('GDT', 'Tổng Cục Thuế – cổng nộp khai thuế etax.gdt.gov.vn.'),
    ('KBNN', 'Kho Bạc Nhà Nước – nơi nộp tiền thuế hàng tháng.'),
]
make_table(
    ['Thuật ngữ', 'Định nghĩa'],
    [(t, d) for t, d in terms],
    col_widths=[4, 12]
)

add_space(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('─' * 60)
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Hathyo MLAM – Claude3 Solution Document')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True; r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Ngày tạo: 17/05/2026 | Phiên bản: 3.0 | Phân tích bởi Claude (Anthropic)')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Thực hiện: Nguyễn Quốc Trọng + Nguyễn Minh Duy | admin@hathyo.com')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.color.rgb = GRAY

# ── Lưu file ──────────────────────────────────────────────────────────────
output_path = r'D:\Hathyo\Hathyo-2026\A1-Phung\analysis\Claude3\ClaudeSolution.docx'
doc.save(output_path)
print(f'OK: {output_path}')
